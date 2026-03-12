from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── NAME ──
p = doc.add_paragraph()
run = p.add_run('Vincent VOUGA')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
p.paragraph_format.space_after = Pt(4)

# ── SUBTITLE ──
p = doc.add_paragraph()
run = p.add_run('Senior Independent IT Consultant | Java & DevOps Specialist')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(4)

# ── CONTACT ──
p = doc.add_paragraph()
run = p.add_run('Brussels, Belgium  |  +32 473 833 209  |  vincent.vouga@gmail.com')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('linkedin.com/in/vincent-vouga  |  github.com/V0OgZ')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('French / Swiss  |  English (C2) · French (Native) · Spanish (B1)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
p.paragraph_format.space_after = Pt(8)

# ── SUMMARY ──
p = doc.add_paragraph()
run = p.add_run('25+ years delivering mission-critical systems in finance, defense, and EU institutions. Java architect with deep expertise in trading systems (FIX, OMS), Spring Boot migration & modernization, DevOps/Cloud transformation, and AI integration. Proven track record leading teams and driving complex projects from design to production in high-stakes, regulated environments.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
run.italic = True
p.paragraph_format.space_after = Pt(12)


def add_section_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    pBdr = p.paragraph_format.element.get_or_add_pPr()
    bottom = pBdr.makeelement(qn('w:pBdr'), {})
    b = bottom.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC'
    })
    bottom.append(b)
    pBdr.append(bottom)
    return p


def add_job(title, company, dates, location, context=None, bullets=None, tags=None):
    # Title line
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    run = p.add_run(f' — {company}')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(1)

    # Dates
    p = doc.add_paragraph()
    run = p.add_run(f'{dates}  |  {location}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(2)

    if context:
        p = doc.add_paragraph()
        run = p.add_run(context)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_after = Pt(2)

    if bullets:
        for b in bullets:
            p = doc.add_paragraph(style='List Bullet')
            if isinstance(b, tuple):
                run = p.add_run(b[0])
                run.bold = True
                run.font.size = Pt(10)
                run = p.add_run(b[1])
                run.font.size = Pt(10)
            else:
                run = p.add_run(b)
                run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)

    if tags:
        p = doc.add_paragraph()
        run = p.add_run(tags)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_after = Pt(8)


# ── TECHNICAL SKILLS ──
add_section_heading('Technical Skills')

skills = [
    ('Backend', 'Java 8→21, Spring Boot 2/3, Microservices, REST, WebSockets'),
    ('DevOps', 'Docker, Kubernetes, Terraform, AWS, Azure, GitHub Actions'),
    ('Data', 'Kafka, Spark, PostgreSQL, Oracle, MongoDB, Elasticsearch'),
    ('Tools', 'Jenkins, Bamboo, Git, Jira, Maven, Nexus'),
    ('AI', 'Generative AI, Agentic programming, Embeddings, Cursor'),
    ('Methods', 'Agile, SAFe, ITIL, Scrum (CSPO®)'),
]

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (cat, desc) in enumerate(skills):
    row = i // 2
    col = i % 2
    cell = table.cell(row, col)
    p = cell.paragraphs[0]
    run = p.add_run(cat.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.add_run('\n')
    run = p.add_run(desc)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── GOVERNMENT & DEFENSE ──
add_section_heading('Government & Defense (2022 – Present) — 3+ Years')

add_job(
    'Senior Software Engineer', 'ILIAS Solutions (NATO Contractor)',
    'May 2024 – Present', 'Brussels, Belgium',
    context='Advanced predictive logistics software for defense environments',
    bullets=[
        ('Spring Boot migration & modernization: ', 'drove migration of legacy modules to Spring Boot 3 / Java 21, restructuring service layers and dependency injection across the microservices ecosystem'),
        'Backend development with Java 21 / Spring Boot across distributed microservices',
        ('Multi-agent programming: ', 'autonomous planning & simulation flows'),
        'AI-assisted development with Cursor to accelerate prototyping, refactoring, and documentation',
        'Local LLM integration in Java (embeddings + lightweight inference for on-premise constraints)',
        'Data pipelines and processing with Kafka / Spark; PostgreSQL persistence',
        'Containerized runtime with Docker Swarm on Linux',
        'Targeted contributions to React/TypeScript front-end',
    ],
    tags='Java 21 • Spring Boot 3 • Kafka • Spark • PostgreSQL • Docker Swarm • Linux • Agentic AI • React/TS'
)

add_job(
    'Application Engineer', 'Council of the European Union',
    'Apr 2022 – Apr 2024', 'Brussels, Belgium',
    context='DevOps Middleware and Service Delivery Engineer (Independent Contractor)',
    bullets=[
        'Application deployment, updates, releases, patches across AWS, Azure, Docker, Kubernetes',
        'Level 2 support; infrastructure decision support for development teams',
        'Maintenance/upgrade of SAP VMs and servers',
        'Automation and operational tooling (Bash / PowerShell)',
        'Runtime/platform operations: Tomcat, IIS; collaboration with delivery teams',
    ],
    tags='Kubernetes • Docker • Terraform • AWS • Azure • Bamboo • Tomcat • PowerShell'
)

# ── FINANCE & TRADING ──
add_section_heading('Finance & Trading Consulting (2006 – 2021) — 15 Years')

add_job(
    'Technology Leader', 'Cream Consulting',
    'Dec 2018 – Nov 2021', 'Brussels',
    bullets=[
        'Led backend development: CRM with Java 8 / Spring Boot, MongoDB, Elasticsearch',
        'Zabbix monitoring; AWS EC2; Agile/Scrum delivery; team mentoring',
    ],
    tags='Java 8 • Spring Boot • MongoDB • Elasticsearch • AWS EC2'
)

add_job(
    'Lead Business Application Support', 'Candriam',
    'Dec 2016 – Nov 2018', 'Paris & Brussels',
    bullets=[
        'L2/L3 support for asset manager (~€100B AUM); led bilingual team of 5',
        'Won service bid for hybrid finance/IT support team; structured escalation and KPIs',
    ],
    tags='Sophis Fusion • Charles River IMS • Oracle • Java 8'
)

add_job(
    'FIX Connectivity Expert', 'BNP Paribas Dealing Services',
    'Sep 2015 – Aug 2016', 'Paris',
    bullets=[
        'FIX protocol certifications; coordinated onboarding with brokers, hubs, and execution venues',
    ],
    tags='FIX 4.2/4.4 • Oracle SQL • Unix'
)

add_job(
    'Java Technical Lead', 'Crédit Agricole CIB',
    'Jul 2014 – Jun 2015', 'Paris',
    bullets=[
        'Java technical reference; WebSphere migration strategy; release governance',
    ],
    tags='Java • WebSphere • Spring • Maven'
)

add_job(
    'Forex Onboarding & Release Manager', 'Société Générale CIB',
    'Jul 2013 – Jun 2014', 'Paris',
    bullets=['FIX APIs for Forex; release management; KPI analysis'],
    tags='Java • Tibco • RMDS'
)

add_job(
    'Java Development Engineer', 'BNP Paribas CIB',
    'Jun 2010 – May 2013', 'Paris',
    bullets=['Developed FIX order router for listed derivatives OMS'],
    tags='Java • FIX • QuickFIX • CORBA • Tibco'
)

add_job(
    'Java Software Engineer', 'ProCapital',
    'Jul 2008 – Jun 2010', 'Paris',
    bullets=['OMS development (SPTrade); FIX messaging; Eclipse RCP'],
    tags='Java • FIX • Eclipse RCP • SonicMQ'
)

add_job(
    'Java Engineer', 'Société Générale CIB',
    'Jul 2006 – Jun 2008', 'Paris',
    bullets=['Equity-derivatives middle-office; CORBA integration'],
    tags='Java • CORBA • Sybase • Tibco'
)

# ── STARTUPS ──
add_section_heading('Startups & Engineering (2000 – 2006)')

add_job(
    'Independent Developer', 'RAINET',
    'Apr 2005 – Jun 2006', 'Paris',
    bullets=['Created Shiva — cybercafé management software, ~40% French market share'],
    tags='Delphi • Interbase'
)

add_job(
    'Software Engineer', 'MOBIQUID (Hi-Media)',
    'Jun 2004 – Mar 2005', 'Paris',
    bullets=['Micropayment solutions (Mediapass/AlloPass) for mobile services'],
    tags='C++ • ASP • SQL Server'
)

add_job(
    'Java Software Engineer', 'COBA Technology',
    'Feb 2002 – May 2004', 'Paris',
    bullets=['Built real-time trading platform for short-term swaps with automatic matching engine'],
    tags='Java 1.4 • Swing • Sockets • Oracle'
)

add_job(
    'Software Engineer', '9 Telecom (Cap Gemini)',
    'Feb 2000 – Dec 2001', 'Paris',
    bullets=['Telephony to Internet migration; C/C++ applications interfacing with BSCS'],
    tags='C • C++ • Unix'
)

# ── PERSONAL PROJECT ──
add_section_heading('Personal Project')

p = doc.add_paragraph()
run = p.add_run('Full-Stack Game Application — Multi-Agent AI Development')
run.bold = True
run.font.size = Pt(10.5)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Designed and orchestrated a multi-agent AI system to develop a large-scale application:')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(2)

for b in [
    'Agent coordination: context summarization, task distribution, inter-agent communication',
    'Backend: Java 21, Spring Boot 3.2, WebSocket/SSE — 950+ classes',
    'Performance: procedural map generation (2M+ tiles) in milliseconds via GraalVM native',
    'Frontend: React 18, Three.js, TypeScript — 320+ files; Mobile: React Native, Expo',
    'Python SDK for offline/online API access; CI/CD with GitHub Actions',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(b)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
run = p.add_run('Multi-Agent AI • Java 21 • GraalVM Native • React • Three.js • Python SDK • TypeScript')
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p.paragraph_format.space_after = Pt(8)

# ── EDUCATION ──
add_section_heading('Education & Certifications')

edu = [
    ('2019', 'Certified Scrum Product Owner® (CSPO®)', 'Scrum Alliance, Belgium'),
    ('1997–1999', "Bachelor's Degree — Mathematics & Computer Science", 'University of Rouen, France'),
    ('1996–1997', 'DUT Software Engineering (Génie Logiciel)', 'Université Le Havre Normandie'),
    ('1995', 'Scientific Baccalauréat (Mathematics)', 'Dieppe, France'),
]
for year, title, place in edu:
    p = doc.add_paragraph()
    run = p.add_run(year)
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(f'    {title}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.add_run('\n')
    run = p.add_run(f'          {place}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(4)

# ── INTERESTS ──
add_section_heading('Interests')
p = doc.add_paragraph()
run = p.add_run('Philosophy, science, technology. Generative AI enthusiast. Secure coding practices.')
run.font.size = Pt(10)

# ── FOOTER ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Last updated: March 2026')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p.paragraph_format.space_before = Pt(16)

doc.save('/Volumes/HOT_DEV/workspace/vincent-vouga/CV_Vincent_VOUGA.docx')
print('DONE - CV_Vincent_VOUGA.docx created')
