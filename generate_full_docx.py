from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page setup ──
for s in doc.sections:
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.2)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

# ── Base style ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.15

DARK = RGBColor(0x1a, 0x1a, 0x1a)
MID = RGBColor(0x4b, 0x55, 0x63)
LIGHT = RGBColor(0x6b, 0x72, 0x80)
FAINT = RGBColor(0x9c, 0xa3, 0xaf)
BORDER_COLOR = 'CCCCCC'


def add_bottom_border(paragraph, color=BORDER_COLOR, sz='6'):
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/></w:pBdr>')
    pPr.append(pBdr)


def add_heading_section(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_bottom_border(p)
    return p


def add_run(p, text, bold=False, italic=False, size=None, color=None):
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def add_job_header(title, company, dates, location):
    p = doc.add_paragraph()
    add_run(p, title, bold=True, size=10.5, color=DARK)
    add_run(p, f'  —  {company}', size=10.5, color=MID)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    add_run(p, f'{dates}  |  {location}', size=9, color=LIGHT)
    p.paragraph_format.space_after = Pt(2)


def add_context(text):
    p = doc.add_paragraph()
    add_run(p, text, italic=True, size=9.5, color=MID)
    p.paragraph_format.space_after = Pt(3)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=10, color=DARK)
        add_run(p, text, size=10, color=MID)
    else:
        add_run(p, text, size=10, color=MID)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15


def add_tags(text):
    p = doc.add_paragraph()
    add_run(p, text, size=8.5, color=MID)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)


# ═══════════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('Vincent VOUGA')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = DARK
p.paragraph_format.space_after = Pt(2)
add_bottom_border(p, color='1A1A1A', sz='12')

p = doc.add_paragraph()
add_run(p, 'Senior Independent IT Consultant  |  Java & DevOps Specialist', size=12, color=MID)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
add_run(p, 'Brussels, Belgium   |   +32 473 833 209   |   vincent.vouga@gmail.com', size=10, color=MID)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
add_run(p, 'linkedin.com/in/vincent-vouga   |   github.com/V0OgZ', size=10, color=MID)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
add_run(p, 'French / Swiss   |   English (C2) \u00b7 French (Native) \u00b7 Spanish (B1)', size=9.5, color=LIGHT)
p.paragraph_format.space_after = Pt(10)

# ── SUMMARY ──
p = doc.add_paragraph()
add_run(p, '25+ years delivering mission-critical systems in finance, defense, and EU institutions. Java architect with deep expertise in trading systems (FIX, OMS), Spring Boot migration & modernization, DevOps/Cloud transformation, and AI integration. Proven track record leading teams and driving complex projects from design to production in high-stakes, regulated environments.', italic=True, size=10, color=MID)
p.paragraph_format.space_after = Pt(10)
add_bottom_border(p, color='E5E7EB')


# ═══════════════════════════════════════════════════════════════
# TECHNICAL SKILLS
# ═══════════════════════════════════════════════════════════════
add_heading_section('Technical Skills')

skills = [
    ('BACKEND', 'Java 8\u219221, Spring Boot 2/3, Microservices, REST, WebSockets'),
    ('DEVOPS', 'Docker, Kubernetes, Terraform, AWS, Azure, GitHub Actions'),
    ('DATA', 'Kafka, Spark, PostgreSQL, Oracle, MongoDB, Elasticsearch'),
    ('TOOLS', 'Jenkins, Bamboo, Git, Jira, Maven, Nexus'),
    ('AI', 'Generative AI, Agentic programming, Embeddings, Cursor'),
    ('METHODS', 'Agile, SAFe, ITIL, Scrum (CSPO\u00ae)'),
]

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (cat, desc) in enumerate(skills):
    row_idx, col_idx = i // 2, i % 2
    cell = table.cell(row_idx, col_idx)
    for existing_p in cell.paragraphs:
        existing_p.clear()
    p = cell.paragraphs[0]
    add_run(p, cat, bold=True, size=8.5, color=DARK)
    p.add_run('\n')
    add_run(p, desc, size=9.5, color=MID)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="60" w:type="dxa"/>'
            f'<w:bottom w:w="60" w:type="dxa"/>'
            f'<w:left w:w="100" w:type="dxa"/>'
            f'<w:right w:w="100" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ═══════════════════════════════════════════════════════════════
# GOVERNMENT & DEFENSE
# ═══════════════════════════════════════════════════════════════
add_heading_section('Government & Defense (2022 \u2013 Present) \u2014 3+ Years')

add_job_header('Senior Software Engineer', 'ILIAS Solutions (NATO Contractor)', 'May 2024 \u2013 Present', 'Brussels, Belgium')
add_context('Advanced predictive logistics software for defense environments')
add_bullet('drove migration of legacy modules to Spring Boot 3 / Java 21, restructuring service layers and dependency injection across the microservices ecosystem', bold_prefix='Spring Boot migration & modernization: ')
add_bullet('Backend development with Java 21 / Spring Boot across distributed microservices')
add_bullet('autonomous planning & simulation flows', bold_prefix='Multi-agent programming: ')
add_bullet('Integration with Cursor (internal tooling for orchestrating advanced business and agent workflows)')
add_bullet('Local LLM integration in Java (embeddings + lightweight inference for on-premise constraints)')
add_bullet('Data pipelines and processing with Kafka / Spark')
add_bullet('Containerized runtime with Docker Swarm, Linux; PostgreSQL persistence')
add_bullet('Targeted contributions to React/TypeScript (non-core UI scope)')
add_tags('Java 21 \u2022 Spring Boot 3 \u2022 Kafka \u2022 Spark \u2022 PostgreSQL \u2022 Docker Swarm \u2022 Linux \u2022 React/TS')

add_job_header('Application Engineer', 'Council of the European Union', 'Apr 2022 \u2013 Apr 2024', 'Brussels, Belgium')
add_context('DevOps Middleware and Service Delivery Engineer (Independent Contractor)')
add_bullet('Application deployment, updates, releases, patches, configuration across Linux, Windows, AWS, Azure, Docker, Kubernetes')
add_bullet('Level 2 support, infrastructure decision support for development teams')
add_bullet('Maintenance/upgrade of SAP VMs and servers')
add_bullet('Automation and operational tooling (Bash / PowerShell)')
add_bullet('Runtime/platform operations: Tomcat, IIS; collaboration with delivery teams')
add_tags('Kubernetes \u2022 Docker \u2022 Terraform \u2022 AWS \u2022 Azure \u2022 Bamboo \u2022 Tomcat \u2022 PowerShell')


# ═══════════════════════════════════════════════════════════════
# FINANCE & TRADING
# ═══════════════════════════════════════════════════════════════
add_heading_section('Finance & Trading (2006 \u2013 2021) \u2014 15 Years')

p = doc.add_paragraph()
add_run(p, 'via Smarfi-services / Open Group \u2014 IT Consultant Specializing in Finance', italic=True, size=9, color=LIGHT)
p.paragraph_format.space_after = Pt(6)

# Cream Consulting
add_job_header('Technology Leader', 'Cream Consulting', 'Dec 2018 \u2013 Nov 2021', 'Brussels')
add_bullet('Led backend development and built a CRM with Java 8, integrating REST APIs, MongoDB, Elasticsearch')
add_bullet('Mentored a junior developer on frontend tasks (React) while maintaining backend delivery ownership')
add_bullet('Implemented Zabbix monitoring; managed IT infrastructure enabling secure remote work')
add_bullet('Agile/Scrum delivery with Jira; maintained CI/CD pipeline and production quality standards')
add_bullet('Managed AWS EC2 instances; defined KPIs; administered Jira/Confluence')
add_bullet('Hands-on development of a Chrome extension (JavaScript) and Node.js scripts')
add_tags('Java 8 \u2022 SparkJava \u2022 Spring Boot \u2022 MongoDB \u2022 Elasticsearch \u2022 OAuth2 \u2022 React \u2022 AWS EC2')

# Candriam
add_job_header('Lead Business Application Support', 'Candriam', 'Dec 2016 \u2013 Nov 2018', 'Paris & Brussels')
add_bullet('Managed Level 2 & 3 business support for an asset manager overseeing \u20ac100B AUM')
add_bullet('Led a bilingual team (EN/FR) of five across Paris and Brussels')
add_bullet('Incident management, KPI tracking, and direct support to trading desks and fund managers')
add_bullet('Assisted and won a service bid for a hybrid finance/IT support team; trained juniors (SQL, log analysis, order flow)')
add_tags('Sophis Fusion \u2022 Charles River IMS \u2022 GPMS \u2022 Oracle \u2022 Java 8')

# BNP Paribas DS
add_job_header('FIX Connectivity Expert', 'BNP Paribas Dealing Services', 'Sep 2015 \u2013 Aug 2016', 'Paris')
add_bullet('Oversaw FIX protocol certifications and coordinated with brokers, FIX networks, and platforms')
add_bullet('Managed a team of three in a bilingual setting; developed support procedures')
add_bullet('Conducted QA and regression testing on trading tools')
add_tags('FIX 4.2/4.4 \u2022 Oracle SQL \u2022 Unix \u2022 QA/Regression')

# Credit Agricole
add_job_header('Java Technical Lead', 'Cr\u00e9dit Agricole CIB', 'Jul 2014 \u2013 Jun 2015', 'Paris')
add_bullet('Technical reference for Java within a service center developing cash management solutions')
add_bullet('Produced technical documentation and standards; facilitated interface between project owners and dev center')
add_bullet('Contributed to WebSphere migration strategy; participated in spec and technical review meetings')
add_tags('Java \u2022 WebSphere \u2022 MQ \u2022 Maven \u2022 Jenkins \u2022 Spring')

# SG CIB Forex
add_job_header('Forex Onboarding & Release Manager', 'Soci\u00e9t\u00e9 G\u00e9n\u00e9rale CIB', 'Jul 2013 \u2013 Jun 2014', 'Paris')
add_bullet('Managed platforms for testing and client integration (onboarding) for SG CIB FIX APIs (Forex)')
add_bullet('Release and environment management; monitoring integration; KPI analysis for capacity planning')
add_tags('Java \u2022 Tibco \u2022 RMDS \u2022 Verifix')

# BNP CIB
add_job_header('Java Development Engineer', 'BNP Paribas CIB', 'Jun 2010 \u2013 May 2013', 'Paris')
add_bullet('Participated in the development of a FIX order router for BNP Paribas brokerage services (listed derivatives OMS)')
add_bullet('Led cross-functional studies to ensure compatibility across processing chains')
add_bullet('Contributed to technical & business improvements; performed FIX certifications with external clients')
add_tags('Java 5 \u2022 FIX \u2022 QuickFIX \u2022 Oracle \u2022 Tibco \u2022 CORBA')

# ProCapital
add_job_header('Java Software Engineer', 'ProCapital', 'Jul 2008 \u2013 Jun 2010', 'Paris')
add_bullet('Developed features for ProCapital\u2019s OMS (SPTrade), deployed to clients for stock market access services')
add_bullet('GUI modifications, FIX message enhancements, and design of a regulatory-compliant data archiving project')
add_tags('Java \u2022 Oracle \u2022 Eclipse RCP \u2022 FIX \u2022 SonicMQ')

# SG CIB 2006
add_job_header('Java Engineer', 'Soci\u00e9t\u00e9 G\u00e9n\u00e9rale CIB', 'Jul 2006 \u2013 Jun 2008', 'Paris')
add_bullet('Developed Java services for an equity-derivatives middle-office repository (Eliot)')
add_bullet('Integrated with surrounding systems using CORBA; maintained and evolved EnorM notification application')
add_tags('Java \u2022 CORBA \u2022 Sybase \u2022 Unix \u2022 Tibco')


# ═══════════════════════════════════════════════════════════════
# STARTUPS
# ═══════════════════════════════════════════════════════════════
add_heading_section('Startups & Engineering (2000 \u2013 2006) \u2014 6 Years')

p = doc.add_paragraph()
add_run(p, 'via CapGemini then independent', italic=True, size=9, color=LIGHT)
p.paragraph_format.space_after = Pt(6)

add_job_header('Independent Developer', 'RAINET', 'Apr 2005 \u2013 Jun 2006', 'Paris')
add_bullet('Created Shiva, a management software suite for cybercaf\u00e9s and multimedia spaces')
add_bullet('Achieved ~40% market share in France')
add_bullet('Client-server architecture, security, licensing, customizable modules')
add_tags('Delphi 5 \u2022 Interbase \u2022 Windows')

add_job_header('Software Engineer', 'MOBIQUID (Hi-Media)', 'Jun 2004 \u2013 Mar 2005', 'Paris')
add_bullet('Developed online mobile services, including micropayment solutions (Mediapass/AlloPass)')
add_tags('C++ \u2022 ASP \u2022 SQL Server')

add_job_header('Java Software Engineer', 'COBA Technology', 'Feb 2002 \u2013 May 2004', 'Paris')
add_bullet('Startup environment: built a real-time trading platform for short-term swaps with automatic matching')
add_bullet('Led discussions with trading specialists, designed a prototype, developed real-time order matching')
add_tags('Java 1.4 \u2022 Swing \u2022 Sockets \u2022 Oracle \u2022 Tomcat')

add_job_header('Software Engineer', '9 Telecom (Cap Gemini)', 'Feb 2000 \u2013 Dec 2001', 'Paris')
add_bullet('Migration of fixed-line telephony services to the Internet; internal applications in C/C++ interfacing with BSCS')
add_tags('C \u2022 C++ \u2022 Windows \u2022 Unix')


# ═══════════════════════════════════════════════════════════════
# PERSONAL PROJECT
# ═══════════════════════════════════════════════════════════════
add_heading_section('Personal Project')

p = doc.add_paragraph()
add_run(p, 'Full-Stack Game Application \u2014 Multi-Agent AI Development', bold=True, size=10.5, color=DARK)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
add_run(p, 'Designed and orchestrated a multi-agent AI system to develop a large-scale application:', size=10, color=MID)
p.paragraph_format.space_after = Pt(3)

add_bullet('Architected agent coordination: context summarization, task distribution, inter-agent communication protocols')
add_bullet('Backend: Java 21, Spring Boot 3.2, WebSocket/SSE \u2014 950+ classes')
add_bullet('Performance: procedural map generation (2M+ tiles) in milliseconds using GraalVM native compilation')
add_bullet('Python SDK: simplified API access for offline/online modes')
add_bullet('Frontend: React 18, Three.js, TypeScript \u2014 320+ files')
add_bullet('Mobile: React Native, Expo')
add_tags('Multi-Agent AI \u2022 Java 21 \u2022 GraalVM Native \u2022 Python SDK \u2022 React \u2022 Three.js')


# ═══════════════════════════════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════════════════════════════
add_heading_section('Education & Certifications')

edu_items = [
    ('2019', 'Certified Scrum Product Owner\u00ae (CSPO\u00ae)', 'Scrum Alliance, Belgium'),
    ('1997\u20131999', 'Bachelor\u2019s Degree \u2014 Mathematics & Computer Science', 'University of Rouen, France'),
    ('1996\u20131997', 'DUT Software Engineering (G\u00e9nie Logiciel)', 'Universit\u00e9 Le Havre Normandie'),
    ('1995', 'Scientific Baccalaur\u00e9at (Mathematics)', 'Dieppe, France'),
]
for year, title, place in edu_items:
    p = doc.add_paragraph()
    add_run(p, f'{year}     ', bold=True, size=10, color=DARK)
    add_run(p, title, size=10, color=DARK)
    p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    add_run(p, f'              {place}', size=9, color=LIGHT)
    p.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════
# INTERESTS
# ═══════════════════════════════════════════════════════════════
add_heading_section('Interests')
p = doc.add_paragraph()
add_run(p, 'Philosophy, science, technology. Generative AI enthusiast. Secure coding practices.', size=10, color=MID)
p.paragraph_format.space_after = Pt(12)

# ── FOOTER ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Last updated: March 2026', size=8, color=FAINT)
p.paragraph_format.space_before = Pt(12)
add_bottom_border(p, color='E5E7EB', sz='4')

OUTPUT = '/Volumes/HOT_DEV/workspace/vincent-vouga/CV_Vincent_VOUGA_Full.docx'
doc.save(OUTPUT)
print(f'DONE: {OUTPUT}')
